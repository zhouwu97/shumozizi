"""对 Pandoc 生成的 DOCX 执行结构检查，并尽力生成可视化复核产物。"""

from __future__ import annotations

import re
import shutil
import subprocess
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader

from shumozizi.core.io import atomic_json, sha256_file

_DOCX_NUMBER = re.compile(r"(?<![0-9A-Za-z.])([+-]?(?:\d+(?:\.\d*)?|\.\d+)(?:[eE][+-]?\d+)?)")


def audit_docx(run_dir: Path, docx_path: Path, *, timeout_seconds: int = 120) -> dict[str, Any]:
    """审计 DOCX 的 ZIP、正文、标题、表图和占位符，并写入 ``qa/docx-structure.json``。

    渲染依赖 LibreOffice，因此它是可选增强；无渲染器时明确记录告警，不把可选 Word
    交付降级为无法冻结的 PDF 编译失败。
    """
    errors: list[str] = []
    warnings: list[str] = []
    zip_members: set[str] = set()
    try:
        with zipfile.ZipFile(docx_path) as archive:
            bad_member = archive.testzip()
            zip_members = set(archive.namelist())
        if bad_member:
            errors.append(f"DOCX ZIP 损坏: {bad_member}")
        if "word/document.xml" not in zip_members:
            errors.append("DOCX 缺少 word/document.xml")
    except (OSError, zipfile.BadZipFile) as exc:
        errors.append(f"DOCX 不是可读取的 ZIP: {exc}")

    paragraphs: list[str] = []
    heading_texts: list[str] = []
    tables = 0
    if not errors:
        try:
            document = Document(docx_path)
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            heading_texts = [
                paragraph.text.strip()
                for paragraph in document.paragraphs
                if paragraph.text.strip() and paragraph.style.name.lower().startswith(("heading", "标题"))
            ]
            tables = len(document.tables)
        except Exception as exc:  # python-docx 会给出底层 XML 读取错误。
            errors.append(f"python-docx 无法读取: {exc}")

    expected_questions: list[str] = []
    state_path = run_dir / "state" / "run.json"
    if state_path.is_file():
        from shumozizi.core.io import load_json

        expected_questions = list(load_json(state_path).get("required_questions", []))
    text = "\n".join(paragraphs)
    missing_questions = [qid for qid in expected_questions if qid not in text]
    if missing_questions:
        errors.append("DOCX 缺少必答问题标题: " + ", ".join(missing_questions))
    unheaded_questions = [qid for qid in expected_questions if not any(qid in heading for heading in heading_texts)]
    if unheaded_questions:
        errors.append("DOCX 必答问题未使用标题层级: " + ", ".join(unheaded_questions))
    placeholders = [token for token in ("TODO", "待填写", "[Paper Title]", "中文摘要内容") if token in text]
    if placeholders:
        errors.append("DOCX 存在占位符: " + ", ".join(placeholders))

    key_metric_ids: list[str] = []
    missing_key_metric_ids: list[str] = []
    try:
        from scripts.qa.metric_ledger import read_ledger
        from shumozizi.simple.results import read_result_index

        ledger = read_ledger(run_dir)
        if ledger:
            current = {
                item["result_id"]: item["metrics"]
                for item in read_result_index(run_dir)["results"]
                if item["status"] == "current"
            }
            displayed_numbers = [(float(match.group(1)), match.group(1)) for match in _DOCX_NUMBER.finditer(text)]
            for metric in ledger["metrics"]:
                if not metric["central"]:
                    continue
                expected = current.get(metric["source_result_id"], {}).get(metric["source_metric"])
                if not isinstance(expected, (int, float)):
                    continue
                key_metric_ids.append(metric["metric_id"])
                if not any(
                    abs(value - float(expected)) <= 0.5 * 10 ** (-len(literal.partition(".")[2]))
                    for value, literal in displayed_numbers
                ):
                    missing_key_metric_ids.append(metric["metric_id"])
            if missing_key_metric_ids:
                errors.append("DOCX 缺少核心指标: " + ", ".join(missing_key_metric_ids))
    except (OSError, ValueError):
        warnings.append("未能读取结果账本，跳过 DOCX 核心指标存在性检查")

    render_pdf: str | None = None
    contact_sheet: str | None = None
    render_page_count: int | None = None
    render_pdf_sha256: str | None = None
    office = (
        shutil.which("soffice.exe")
        or shutil.which("soffice")
        or shutil.which("libreoffice.exe")
        or shutil.which("libreoffice")
    )
    if office:
        qa_dir = run_dir / "qa"
        qa_dir.mkdir(parents=True, exist_ok=True)
        try:
            completed = subprocess.run(
                [office, "--headless", "--convert-to", "pdf", "--outdir", str(qa_dir), str(docx_path)],
                capture_output=True,
                text=True,
                timeout=timeout_seconds,
                check=False,
            )
            candidate = qa_dir / f"{docx_path.stem}.pdf"
            target = qa_dir / "docx-render.pdf"
            if completed.returncode == 0 and candidate.is_file() and candidate.stat().st_size > 0:
                candidate.replace(target)
                render_pdf = str(target.relative_to(run_dir))
                render_pdf_sha256 = sha256_file(target)
                try:
                    render_page_count = len(PdfReader(str(target)).pages)
                except Exception as exc:  # pypdf 的损坏 PDF 异常随版本变化。
                    warnings.append(f"DOCX 渲染 PDF 页数读取失败: {exc}")
                try:
                    from tools.qa.make_contact_sheet import make_contact_sheet

                    sheet = qa_dir / "docx-contact-sheet.png"
                    make_contact_sheet(target, sheet)
                    contact_sheet = str(sheet.relative_to(run_dir))
                except Exception as exc:  # 联系表只服务人工复核，不能掩盖已完成的结构检查。
                    warnings.append(f"DOCX 联系表未生成: {exc}")
            else:
                warnings.append("LibreOffice 未能生成 DOCX 渲染 PDF")
        except (OSError, subprocess.TimeoutExpired) as exc:
            warnings.append(f"DOCX 渲染跳过: {exc}")
    else:
        warnings.append("未检测到 LibreOffice，未生成 DOCX 渲染 PDF/联系表")

    report = {
        "schema_version": "1.0",
        "docx_path": str(docx_path.relative_to(run_dir)),
        "success": not errors,
        "errors": errors,
        "warnings": warnings,
        "paragraph_count": len(paragraphs),
        "table_count": tables,
        "image_count": len([name for name in zip_members if name.startswith("word/media/")]),
        "missing_question_titles": missing_questions,
        "unheaded_question_titles": unheaded_questions,
        "placeholder_count": len(placeholders),
        "key_metric_ids": key_metric_ids,
        "missing_key_metric_ids": missing_key_metric_ids,
        "render_pdf": render_pdf,
        "render_pdf_sha256": render_pdf_sha256,
        "render_page_count": render_page_count,
        "contact_sheet": contact_sheet,
    }
    atomic_json(run_dir / "qa" / "docx-structure.json", report)
    return report
