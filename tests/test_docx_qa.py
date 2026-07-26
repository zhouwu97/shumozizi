"""DOCX 结构化 QA 回归。"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from shumozizi.paper.docx_qa import audit_docx
from shumozizi.simple.initialization import initialize_simple_run


def test_docx_qa_reads_document_and_required_question_title(tmp_path: Path) -> None:
    """可读 Word、Q 标题和表格计数必须进入结构报告。"""
    run_dir = initialize_simple_run(tmp_path, "docx-qa", required_questions=["Q1"])
    path = run_dir / "paper" / "final.docx"
    document = Document()
    document.add_heading("问题 Q1", level=1)
    document.add_paragraph("正文保留关键结论。")
    document.add_table(rows=1, cols=1).cell(0, 0).text = "指标"
    document.save(path)

    report = audit_docx(run_dir, path)

    assert report["success"] is True
    assert report["table_count"] == 1
    assert report["missing_question_titles"] == []
    assert (run_dir / "qa" / "docx-structure.json").is_file()


def test_docx_qa_rejects_non_zip_file(tmp_path: Path) -> None:
    """伪造扩展名的非 ZIP 文件不能作为已验证 DOCX 交付。"""
    run_dir = initialize_simple_run(tmp_path, "broken-docx")
    path = run_dir / "paper" / "final.docx"
    path.write_bytes(b"not-a-docx")

    report = audit_docx(run_dir, path)

    assert report["success"] is False
    assert any("ZIP" in error for error in report["errors"])
